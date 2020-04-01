#! python3

import docx, os
from docx.enum.text import WD_ALIGN_PARAGRAPH

guests = open('guests.txt')
doc = docx.Document()
for guest in guests.readlines():
	guest = guest.strip()

	para0 = doc.add_paragraph('It would be a pleasure to have the company of')
	para0.style = doc.styles['Custom 1']
	para0.alignment = WD_ALIGN_PARAGRAPH.CENTER

	name = doc.add_paragraph(guest)
	name.alignment = WD_ALIGN_PARAGRAPH.CENTER

	para2 = doc.add_paragraph('at 110101 Memory Lane on the Evening of')
	para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

	date = doc.add_paragraph('April 1st')
	date.alignment = WD_ALIGN_PARAGRAPH.CENTER

	time = doc.add_paragraph("at 7pm")
	time.alignment = WD_ALIGN_PARAGRAPH.CENTER

	doc.add_page_break()

doc.save('sample.docx')
